import streamlit as st
import anthropic
import requests
import os
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
from datetime import datetime
from io import BytesIO
from pathlib import Path

load_dotenv()

# ── API KEYS ──────────────────────────────────────────────────────────────────
try:
    api_key = st.secrets["ANTHROPIC_API_KEY"]
    UNSPLASH_KEY = st.secrets["UNSPLASH_ACCESS_KEY"]
except:
    api_key = os.getenv("ANTHROPIC_API_KEY")
    UNSPLASH_KEY = os.getenv("UNSPLASH_ACCESS_KEY")

client = anthropic.Anthropic(api_key=api_key)

# ── PAGE CONFIG ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Editorial AI Studio — Industrial Tech",
    page_icon="⚙️",
    layout="wide"
)

st.markdown("""
<style>
    .stApp { background-color: #0d0d0d; }
    .stButton > button { background-color: #f5c400; color: #000; border: none;
        border-radius: 4px; font-weight: bold; }
    .stButton > button:hover { background-color: #d4a900; }
    h2, h3 { color: #f5c400 !important; }
    .stSelectbox label, .stMultiSelect label, .stTextInput label,
    .stTextArea label, .stSlider label { color: #ffffff !important; }
    p, .stMarkdown { color: #dddddd; }
</style>
""", unsafe_allow_html=True)

# ── HEADER ────────────────────────────────────────────────────────────────────
col_logo, col_title = st.columns([0.3, 0.7])
with col_logo:
    logo_path = Path(__file__).parent / "logo.png"
    if logo_path.exists():
        st.image(str(logo_path), width=300)
    else:
        st.markdown("### ⚙️ Industrial Tech")
with col_title:
    st.markdown("<h2 style='color:#f5c400;margin-top:20px'>Editorial AI Studio</h2>", unsafe_allow_html=True)
    st.markdown("<span style='color:#aaaaaa'>Sistema editoriale intelligente — powered by Claude AI</span>", unsafe_allow_html=True)

st.divider()

# ── SESSION STATE ─────────────────────────────────────────────────────────────
for key in ["briefing", "temi_proposti", "temi_selezionati", "articoli_generati"]:
    if key not in st.session_state:
        st.session_state[key] = {} if key == "briefing" else []

# ── MODULO 1: BRIEFING ────────────────────────────────────────────────────────
st.markdown("## 📋 Modulo 1 — Briefing del fascicolo")

with st.form("briefing_form"):
    col1, col2 = st.columns(2)
    with col1:
        numero = st.text_input("Numero / Titolo fascicolo", placeholder="Es: Speciale Hannover Messe 2026")
        lingua = st.selectbox("Lingua principale", ["Italiano", "Inglese", "Bilingue (IT/EN)"])
        settori = st.multiselect("Settori tematici", [
            "Automazione industriale", "Oil & Gas", "Energia & Utilities",
            "Robotica", "IIoT / Industria 4.0", "Manutenzione predittiva",
            "Sicurezza industriale", "Procurement & Supply chain",
            "Sostenibilità & Green Energy", "Digitale & AI"
        ])
        fiere = st.text_input("Fiere / eventi collegati", placeholder="Es: Hannover Messe, SPS, Adipec")
    with col2:
        paesi = st.multiselect("Paesi di distribuzione", [
            "Italia", "Germania", "UK", "Francia", "Spagna",
            "USA", "Middle East", "Asia Pacific", "Benelux"
        ])
        inserzionisti = st.text_area("Inserzionisti target / aziende da citare",
            placeholder="Es: Siemens, ABB, Emerson, Honeywell, Rockwell...", height=80)
        note = st.text_area("Note libere per la redazione",
            placeholder="Angolazione editoriale, temi da evitare, messaggi chiave...", height=80)

    n_articoli = st.slider("Numero articoli originali", 3, 8, 5)
    n_comunicati = st.slider("Numero comunicati da rielaborare", 5, 15, 10)
    submitted = st.form_submit_button("💾 Salva briefing e proponi temi →")

if submitted:
    st.session_state.briefing = {
        "numero": numero, "lingua": lingua, "settori": settori,
        "fiere": fiere, "paesi": paesi, "inserzionisti": inserzionisti,
        "note": note, "n_articoli": n_articoli, "n_comunicati": n_comunicati
    }
    st.success(f"✅ Briefing salvato per: **{numero}**")

# ── MODULO 2: ARTICOLI ────────────────────────────────────────────────────────
if st.session_state.briefing:
    st.divider()
    st.markdown("## ✍️ Modulo 2 — Proposta e generazione articoli")
    b = st.session_state.briefing

    if st.button("🧠 Genera proposta temi articoli"):
        with st.spinner("Claude sta elaborando i temi per questo fascicolo..."):
            prompt = f"""Sei il direttore editoriale di Industrial Tech Magazine, rivista B2B italiana specializzata in automazione industriale e Oil & Gas, distribuita anche internazionalmente.

Fascicolo: {b['numero']}
Lingua: {b['lingua']}
Settori: {', '.join(b['settori'])}
Fiere/eventi: {b['fiere']}
Paesi distribuzione: {', '.join(b['paesi'])}
Aziende/inserzionisti da considerare: {b['inserzionisti']}
Note redazionali: {b['note']}

Proponi {b['n_articoli'] + 3} idee di articoli originali per questo fascicolo.
Per ognuno fornisci:
1. TITOLO: titolo giornalistico accattivante
2. OCCHIELLO: sottotitolo di una riga
3. TEMA: tema principale in 2 righe
4. ANGOLO: angolazione editoriale unica (perché è rilevante ADESSO)
5. FONTI: 2-3 trend/dati/eventi reali su cui basarsi
6. LINGUA: IT o EN

Rispondi SOLO con un JSON array con questi campi: titolo, occhiello, tema, angolo, fonti, lingua"""

            response = client.messages.create(
                model="claude-sonnet-4-5",
                max_tokens=3000,
                messages=[{"role": "user", "content": prompt}]
            )
            raw = response.content[0].text
            try:
                start = raw.find('[')
                end = raw.rfind(']') + 1
                st.session_state.temi_proposti = json.loads(raw[start:end])
            except:
                st.error("Errore nel parsing JSON. Riprova.")

    if st.session_state.temi_proposti:
        st.markdown("### Seleziona gli articoli da produrre:")
        selezionati = []
        for i, t in enumerate(st.session_state.temi_proposti):
            col1, col2 = st.columns([0.08, 0.92])
            with col1:
                sel = st.checkbox("", key=f"tema_{i}")
            with col2:
                with st.expander(f"**{t.get('titolo', '')}** — {t.get('occhiello', '')}"):
                    st.markdown(f"**Tema:** {t.get('tema', '')}")
                    st.markdown(f"**Angolo:** {t.get('angolo', '')}")
                    st.markdown(f"**Fonti:** {t.get('fonti', '')}")
                    st.markdown(f"**Lingua:** {t.get('lingua', 'IT')}")
            if sel:
                selezionati.append(t)

        st.session_state.temi_selezionati = selezionati
        st.info(f"Articoli selezionati: **{len(selezionati)}**")

        if selezionati and st.button("📝 Genera articoli selezionati → .docx"):
            articoli = []
            progress = st.progress(0)

            for idx, tema in enumerate(selezionati):
                with st.spinner(f"Scrivo: {tema['titolo']}..."):
                    lingua_art = "inglese" if tema.get('lingua', 'IT') == 'EN' else "italiano"
                    prompt_art = f"""Sei un giornalista senior di Industrial Tech Magazine. Scrivi un articolo professionale in {lingua_art}.

Titolo: {tema['titolo']}
Occhiello: {tema['occhiello']}
Tema: {tema['tema']}
Angolo: {tema['angolo']}
Riferimenti: {tema['fonti']}
Fascicolo: {st.session_state.briefing['numero']}
Distribuzione: {', '.join(st.session_state.briefing['paesi'])}

ISTRUZIONI:
- Lunghezza: circa 4000 battute (spazi inclusi)
- Struttura: lead forte → 3-4 sezioni con sottotitolo → conclusione con outlook
- Tono: tecnico ma accessibile, autorevole, giornalistico B2B
- Dati, percentuali, citazioni da esperti del settore
- NON usare clichè da AI
- Sottotitoli marcati con ###"""

                    resp = client.messages.create(
                        model="claude-sonnet-4-5",
                        max_tokens=2000,
                        messages=[{"role": "user", "content": prompt_art}]
                    )
                    testo = resp.content[0].text

                    foto_url, foto_credits = "", ""
                    try:
                        query = " ".join(tema['titolo'].split()[:3]) + " industrial"
                        r = requests.get(
                            "https://api.unsplash.com/search/photos",
                            params={"query": query, "per_page": 1, "orientation": "landscape"},
                            headers={"Authorization": f"Client-ID {UNSPLASH_KEY}"}
                        )
                        data = r.json()
                        if data.get("results"):
                            foto_url = data["results"][0]["urls"]["regular"]
                            foto_credits = f"Photo by {data['results'][0]['user']['name']} on Unsplash"
                    except:
                        pass

                    articoli.append({
                        "titolo": tema['titolo'], "occhiello": tema['occhiello'],
                        "testo": testo, "lingua": tema.get('lingua', 'IT'),
                        "foto_url": foto_url, "foto_credits": foto_credits
                    })
                progress.progress((idx + 1) / len(selezionati))

            st.session_state.articoli_generati = articoli

            doc = Document()
            doc.styles['Normal'].font.name = 'Georgia'
            doc.styles['Normal'].font.size = Pt(11)

            title_par = doc.add_paragraph()
            title_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title_par.add_run(f"Industrial Tech Magazine — {st.session_state.briefing['numero']}")
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(245, 196, 0)
            doc.add_paragraph(f"Generato il {datetime.now().strftime('%d/%m/%Y %H:%M')}")
            doc.add_page_break()

            for art in articoli:
                p = doc.add_paragraph()
                r = p.add_run(art['titolo'])
                r.font.size = Pt(18)
                r.font.bold = True
                r.font.color.rgb = RGBColor(245, 196, 0)

                p2 = doc.add_paragraph()
                r2 = p2.add_run(art['occhiello'])
                r2.font.size = Pt(13)
                r2.font.italic = True

                doc.add_paragraph(f"[Lingua: {art['lingua']}]")
                if art['foto_url']:
                    doc.add_paragraph(f"📷 FOTO: {art['foto_url']}")
                    doc.add_paragraph(f"Credits: {art['foto_credits']}")
                doc.add_paragraph("")

                for riga in art['testo'].split('\n'):
                    if riga.startswith('### '):
                        p = doc.add_paragraph()
                        r = p.add_run(riga.replace('### ', ''))
                        r.font.size = Pt(13)
                        r.font.bold = True
                    elif riga.strip():
                        doc.add_paragraph(riga)
                doc.add_page_break()

            buf = BytesIO()
            doc.save(buf)
            buf.seek(0)
            fname = f"articoli_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
            st.download_button("📥 Scarica articoli .docx", buf, file_name=fname,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            st.balloons()

# ── MODULO 3: COMUNICATI ──────────────────────────────────────────────────────
st.divider()
st.markdown("## 📰 Modulo 3 — Rielaborazione comunicati stampa")

comunicati_input = st.text_area(
    "Incolla qui i comunicati stampa (separali con --- su una riga)",
    height=250,
    placeholder="Incolla il testo del comunicato 1...\n---\nIncolla il testo del comunicato 2..."
)

col1, col2 = st.columns(2)
with col1:
    formato = st.selectbox("Formato output", [
        "Breve (mezza pagina ~800 battute)",
        "Medio (una pagina ~1800 battute)",
        "Lungo (due pagine ~3500 battute)"
    ])
with col2:
    tono = st.selectbox("Tono redazionale", [
        "Giornalistico neutro",
        "Tecnico approfondito",
        "Business & management",
        "Innovazione & trend"
    ])

if comunicati_input and st.button("✂️ Rielabora comunicati → .docx"):
    comunicati = [c.strip() for c in comunicati_input.split('---') if c.strip()]
    battute_map = {"Breve": 800, "Medio": 1800, "Lungo": 3500}
    n_batt = battute_map.get(formato.split()[0], 1800)

    doc2 = Document()
    p = doc2.add_paragraph()
    r = p.add_run(f"Comunicati rielaborati — {st.session_state.briefing.get('numero', 'fascicolo')}")
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = RGBColor(245, 196, 0)
    doc2.add_paragraph(f"Formato: {formato} | Tono: {tono}")
    doc2.add_page_break()

    progress2 = st.progress(0)
    for idx, com in enumerate(comunicati):
        with st.spinner(f"Rielaboro comunicato {idx+1}/{len(comunicati)}..."):
            prompt_com = f"""Sei un redattore di Industrial Tech Magazine. Rielabora questo comunicato stampa.

COMUNICATO ORIGINALE:
{com[:3000]}

ISTRUZIONI:
- Lunghezza target: circa {n_batt} battute
- Tono: {tono}
- Trasforma da comunicato aziendale a notizia redazionale
- Aggiungi contesto di settore dove utile
- Titolo accattivante + lead forte
- NON sembrare un comunicato stampa
- Mantieni la lingua originale del comunicato

Rispondi con: TITOLO: ... poi a capo il testo."""

            resp = client.messages.create(
                model="claude-sonnet-4-5",
                max_tokens=1500,
                messages=[{"role": "user", "content": prompt_com}]
            )
            testo_com = resp.content[0].text
            lines = testo_com.strip().split('\n')
            titolo_com = lines[0].replace('TITOLO:', '').strip() if lines else f"Comunicato {idx+1}"
            corpo = '\n'.join(lines[1:]).strip()

            p = doc2.add_paragraph()
            r = p.add_run(titolo_com)
            r.font.size = Pt(15)
            r.font.bold = True
            r.font.color.rgb = RGBColor(245, 196, 0)
            doc2.add_paragraph(f"[{formato} | {tono}]")
            doc2.add_paragraph("")
            for riga in corpo.split('\n'):
                if riga.strip():
                    doc2.add_paragraph(riga)
            doc2.add_page_break()
        progress2.progress((idx + 1) / len(comunicati))

    buf2 = BytesIO()
    doc2.save(buf2)
    buf2.seek(0)
    fname2 = f"comunicati_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    st.download_button("📥 Scarica comunicati .docx", buf2, file_name=fname2,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    st.success("✅ Pronto!")